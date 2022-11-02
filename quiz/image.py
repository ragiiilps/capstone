import requests
import io
from docx import Document
from docx.shared import Inches
from PIL import Image

# url = 'https://firebasestorage.googleapis.com/v0/b/aigyujiem-image.appspot.com/o/images%2F1624186128.8245711.png?alt=media'
# response = requests.get(url, stream=True)
# image = io.BytesIO(response.content)
# image2 = Image.open(image)
# width, height = image2.size
# document = Document()
# p = document.add_paragraph('A plain paragraph having some ')
# document.add_paragraph('Another sentence')
# # if width >= 288 and width >= height:
# #     document.add_picture(image, width=Inches(3.0))
# # elif height >= 192 and height > width:
# #     document.add_picture(image, height=Inches(2.0))

# document.save('demo.docx')

# payload = {
#             "username": "dosen_silmi",
#             "password" : "Capstone_aig123"
#         }

# request = requests.post("https://moodle.aigyujiem.com/login/token.php?service=ugm", params = payload)
# print(request.content)

import requests

url = "https://moodle.aigyujiem.com/login/token.php?username=dosen_silmi&password=Capstone_aig123&service=ugm"

payload={}
headers = {"User-Agent":"Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.97 Safari/537.36"}

response = requests.request("GET", url, headers=headers, data=payload)

print(response.json()['token'])

# import http.client

# conn = http.client.HTTPSConnection("moodle.aigyujiem.com")
# payload = ''
# headers = {}
# conn.request("GET", "/login/token.php?username=dosen_silmi&password=Capstone_aig123&service=ugm", payload, headers)
# res = conn.getresponse()
# data = res.read()
# # print(data.decode("utf-8").split('"')[3])

# payload_userid = {
#                 "wstoken": "5b453df2205a5ee8759a611aacff5bd5",
#                 "criteria[0][key]": "username",
#                 "criteria[0][value]": "dosen_silmi"
#             }


# conn.request("POST", "/webservice/rest/server.php?wsfunction=core_user_get_users&moodlewsrestformat=json", payload_userid)
# res = conn.getresponse()
# data = res.read()
# print(data)
# print(data.decode("utf-8"))
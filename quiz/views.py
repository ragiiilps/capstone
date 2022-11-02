from asyncio.windows_events import NULL
from xmlrpc.client import DateTime
from .helper.kisi import *
from django.shortcuts import render
from django.shortcuts import redirect, get_object_or_404
from django.contrib.auth.decorators import user_passes_test
from django.contrib.auth.models import Group
from django.http import Http404, HttpResponse, HttpResponseRedirect, FileResponse
from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from .form import *
from .helper.parser import *
from .helper.generator import *
from django.http import HttpResponse
from .models import *
from html import unescape
from django.urls import reverse, reverse_lazy
from urllib.parse import urlencode
from django.utils.html import strip_tags, escape
from html import unescape
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime
from bs4 import BeautifulSoup
from lxml import etree
from django.http import JsonResponse
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
from aig.settings import MEDIA_ROOT, MEDIA_URL, STATIC_ROOT, STATIC_URL
from django.core.files.storage import default_storage
from django.contrib import messages
import pyrebase
import latex2mathml.converter as mcv
from ast import literal_eval
import random
import requests
import io
import base64
import logging
import os
from zipfile import ZipFile
from PIL import Image
import logging
from .helper.kisi import *
from .similaritychecker.similaritychecker import check_similarity
from .emailsender.emailsender import send_reviewer_notif, send_writer_notif, send_penulis_notif_reviewed, send_reviewer_notif_new
from django.contrib import messages
from xhtml2pdf import pisa
from django.views.generic import View
from django.template.loader import get_template
import threading
import collections
collections.Callable = collections.abc.Callable


@csrf_exempt
def Index(request):

    # moodle_api = "https://moodle.aigyujiem.com/login/token.php?service=ugm"
    # moodle_api_userid = "https://moodle.aigyujiem.com/webservice/rest/server.php?wsfunction=core_user_get_users&moodlewsrestformat=json"
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.97 Safari/537.36"}

    if request.method == "GET":
        if request.user.is_authenticated:
            return redirect('dashboard')
        else:
            return render(request, 'index.html')

    if request.method == "POST":

        username_login = request.POST['username']
        password_login = request.POST['password']

        payload = {
            "username": username_login,
            "password": password_login
        }

        user = authenticate(request, username=username_login,
                            password=password_login)

        if user is not None:
            login(request, user)
            # moodle_response_token = requests.get(moodle_api, params=payload, headers=headers)
            # request.session['moodle_token'] = moodle_response_token.json()['token']

            # payload_userid = {
            #     "wstoken": request.session['moodle_token'],
            #     "criteria[0][key]": "username",
            #     "criteria[0][value]": username_login
            # }

            # moodle_response_userid = requests.post(moodle_api_userid, params=payload_userid, headers=headers)
            # request.session['user_id'] = moodle_response_userid.json()['users'][0]['id']
            last_url = request.GET.get('url')
            if last_url is not None:
                return redirect(last_url)
            return redirect('dashboard')
        else:
            return redirect('index')


@csrf_exempt
def Dashboard(request):
    if request.user.is_authenticated:
        user_id = request.user.id
        user_profile = Profile.objects.get(id=user_id)
        first_name = user_profile.user.first_name
        last_name = user_profile.user.last_name
        full_name = first_name + ' ' + last_name
        email = user_profile.user.email
        instansi = user_profile.sekolah
        mapel = user_profile.mapel
        jabatan = user_profile.jabatan
        golongan = user_profile.golongan
        alamat = user_profile.alamat
        nip = user_profile.nip
        nomer_hp = user_profile.nomer_hp
        pangkat = user_profile.pangkat

        context = {
            'ceklogin': request.user.is_authenticated,
            'username': request.user,
            'usergroup': request.user.groups.all,
            'testgroup': Group.objects.get(name='creator'),
            'dashboard': 'active',
            'full_name': full_name,
            'email': email,
            'instansi': instansi,
            'mapel': mapel,
            'nip': nip,
            'golongan': golongan,
            'pangkat': pangkat,
            'alamat': alamat,
            'nomer_hp': nomer_hp,
            'jabatan': jabatan,
        }
        return render(request, 'dashboard.html', context)
    else:
        return redirect('index')


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def IdentitasItem(request):
    if request.method == 'POST':
        form = IdentitasSoal(request.POST)
        if form.is_valid():
            indikator_id = request.session.get('indikator_id')
            inus_id = request.session.get('inus_id')
            request.session['indikator_id'] = indikator_id
            request.session['inus_id'] = inus_id
            soal = form.cleaned_data['soal']
            request.session['stem'] = soal
            request.session['amount'] = form.cleaned_data['jum_soal']
            request.session['arrange'] = form.cleaned_data['acak']
            base_url = reverse('prosesitem')
            if 'uji-kualitas' in request.POST:
                penugasan_id = request.session.get('penugasan_id')
                penugasan = Penugasan.objects.get(id=penugasan_id)
                matkul = penugasan.mata_kuliah
                first_closest, second_closest = check_similarity(soal, matkul)
                link_base = reverse('showpdf')
                query_first = urlencode({
                    'matkul': first_closest['matkul'],
                    'file': first_closest['file_name']
                })
                query_second = urlencode({
                    'matkul': second_closest['matkul'],
                    'file': second_closest['file_name']
                })
                # base_url = reverse('identitas')
                return render(request, 'identitas.html', {
                    'identitas': 'active',
                    'uji_kualitas': True,
                    'first_closest': first_closest,
                    'second_closest': second_closest,
                    'first_link': '{}?{}'.format(link_base, query_first),
                    'second_link': '{}?{}'.format(link_base, query_second)
                })

            if request.GET.get('edit'):
                query_string = urlencode({
                    'edit': 'true',
                    'id': request.GET.get('id')
                })
                url = '{}?{}'.format(base_url, query_string)
                return redirect(url)
            elif request.GET.get('kembali'):
                query_string = urlencode({
                    'edit': 'true',
                    'id': request.GET.get('id')
                })
                url = '{}?{}'.format(base_url, query_string)
                return redirect(url)
            else:
                return redirect(base_url)

        else:
            request.session['stem'] = request.POST.get('soal')
            request.session['amount'] = request.POST.get('jum_soal')
            request.session['arrange'] = request.POST.get('acak')
            return render(request, 'identitas.html', {
                'form': form,
                'identitas': 'active'
            })
    elif request.method == 'GET':
        if request.GET.get('fromtemplate'):
            print('template')
            return render(request, 'identitas.html', {
                'identitas': 'active'
            })
        elif request.GET.get('edit'):
            print('edit')
            return render(request, 'identitas.html', {
                'identitas': 'active'
            })
        else:
            request.session['stem'] = ''
            request.session['amount'] = ''
            request.session['arrange'] = ''
            request.session['var_value'] = ''
            request.session['formula'] = ''
            request.session['option'] = ''
            request.session['var_form_value'] = ''
            request.session['question'] = ''
            request.session['answer'] = ''
            request.session['img_var'] = ''
        return render(request, 'identitas.html', {
            'identitas': 'active'
        })


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def ProsesItem(request):
    stem = request.session.get('stem')
    indikator_id = request.session.get('indikator_id')
    inus_id = request.session.get('inus_id')
    var_id, var_name = getValue1(stem)
    var_form_value = gen_var_form(var_name)
    request.session['var_form_value'] = var_form_value

    if str(var_name) != request.session.get('var_name'):
        print('changed')
        print(str(var_name) + ' vs ' + str(request.session.get('var_name')))
        request.session['var_name'] = var_name

        var_form_value = gen_var_form(var_name)
        request.session['var_form_value'] = var_form_value
    else:
        print('unchanged')
        request.session['var_form_value'] = request.session.get('var_value')

    if request.method == 'GET':
        return render(request, 'prosesitem.html')
    elif request.method == 'POST':
        id_soal = request.POST.get('id')
        form = MultiFormula(request.POST, request.FILES)
        print(request.POST.get('variabel'))

        if form.is_valid():
            request.session['indikator_id'] = indikator_id
            request.session['inus_id'] = inus_id
            request.session['var_value'] = form.cleaned_data['variabel']
            request.session['formula'] = form.cleaned_data['multirumus']
            request.session['option'] = unescape(form.cleaned_data['opsi'])
            request.session['img_var'] = form.cleaned_data['img_var']

            if request.GET.get('edit'):
                base_url = reverse('result')
                query_string = urlencode({
                    'edit': 'true',
                    'id': request.GET.get('id')
                })
                url = '{}?{}'.format(base_url, query_string)
                return redirect(url)
            else:
                return redirect('result')
        else:
            print('invalid 174')
            request.session['var_form_value'] = request.POST.get('variabel')
            request.session['formula'] = request.POST.get('multirumus')
            request.session['option'] = request.POST.get('opsi')
            if request.session['img_var']:
                request.session['img_var'] = request.POST.get('img_var')
            return render(request, 'prosesitem.html', {'form': form})


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def BuatTemplate(request):
    if request.method == 'POST':
        form = CreateTemplate(request.POST)
        if form.is_valid():
            indikator_id = request.session.get('indikator_id')
            inus_id = request.session.get('inus_id')
            request.session['indikator_id'] = indikator_id
            request.session['inus_id'] = inus_id
            request.session['stem'] = form.cleaned_data['soal']
            request.session['amount'] = form.cleaned_data['jum_soal']
            request.session['arrange'] = form.cleaned_data['acak']
            if request.GET.get('edit'):
                base_url = reverse('prosestemplate')
                query_string = urlencode({
                    'edit': 'true',
                    'id': request.GET.get('id')
                })
                url = '{}?{}'.format(base_url, query_string)
                return redirect(url)
            elif request.GET.get('kembali'):
                base_url = reverse('prosestemplate')
                query_string = urlencode({
                    'edit': 'true',
                    'id': request.GET.get('id')
                })
                url = '{}?{}'.format(base_url, query_string)
                return redirect(url)
            else:
                return redirect('prosestemplate')

        else:
            request.session['judul'] = request.POST.get('judul')
            request.session['stem'] = request.POST.get('soal')
            request.session['amount'] = request.POST.get('jum_soal')
            request.session['arrange'] = request.POST.get('acak')
            return render(request, 'buat_template.html', {
                'form': form,
                'buattemplate': 'active'
            })
    elif request.method == 'GET':
        if request.GET.get('fromtemplate'):
            print('template')
            return render(request, 'buat_template.html', {
                'buattemplate': 'active'
            })
        elif request.GET.get('edit'):
            print('edit')
            return render(request, 'buat_template.html', {
                'buattemplate': 'active'
            })
        else:
            request.session['judul'] = ''
            request.session['stem'] = ''
            request.session['amount'] = ''
            request.session['arrange'] = ''
            request.session['var_value'] = ''
            request.session['formula'] = ''
            request.session['option'] = ''
            request.session['var_form_value'] = ''
            request.session['question'] = ''
            request.session['answer'] = ''
            request.session['img_var'] = ''
            return render(request, 'buat_template.html', {
                'buattemplate': 'active'
            })


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def ProsesTemplate(request):
    stem = request.session.get('stem')
    indikator_id = request.session.get('indikator_id')
    inus_id = request.session.get('inus_id')
    var_id, var_name = getValue1(stem)
    var_form_value = gen_var_form(var_name)
    request.session['var_form_value'] = var_form_value

    if str(var_name) != request.session.get('var_name'):
        print('changed')
        print(str(var_name) + ' vs ' + str(request.session.get('var_name')))
        request.session['var_name'] = var_name

        var_form_value = gen_var_form(var_name)
        request.session['var_form_value'] = var_form_value
    else:
        print('unchanged')
        request.session['var_form_value'] = request.session.get('var_value')

    if request.method == 'GET':
        return render(request, 'prosestemplate.html')
    elif request.method == 'POST':
        id_soal = request.POST.get('id')
        form = MultiFormula(request.POST)
        print(form)
        print(request.POST.get('variabel'))

        if form.is_valid():
            request.session['indikator_id'] = indikator_id
            request.session['inus_id'] = inus_id
            request.session['var_value'] = form.cleaned_data['variabel']
            request.session['formula'] = form.cleaned_data['multirumus']
            request.session['option'] = unescape(form.cleaned_data['opsi'])
            request.session['img_var'] = form.cleaned_data['img_var']

            if request.GET.get('edit'):
                base_url = reverse('result_template')
                query_string = urlencode({
                    'edit': 'true',
                    'id': request.GET.get('id')
                })
                url = '{}?{}'.format(base_url, query_string)
                return redirect(url)
            else:
                return redirect('result_template')
        else:
            print('invalid hiyak')
            request.session['var_form_value'] = request.POST.get('variabel')
            request.session['formula'] = request.POST.get('multirumus')
            request.session['option'] = request.POST.get('opsi')
            if request.session['img_var']:
                request.session['img_var'] = request.POST.get('img_var')
            return render(request, 'prosestemplate.html', {'form': form})


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def ResultTemplate(request):
    indikator_id = request.session.get('indikator_id')
    inus_id = request.session.get('inus_id')
    id_creator = request.user.id
    stem = request.session.get('stem')
    amount = request.session.get('amount')
    arrange = request.session.get('arrange')
    var_value = request.session.get('var_value')
    var_name = request.session.get('var_name')
    formula = request.session.get('formula')
    option = request.session.get('option')
    img_var = request.session.get('img_var')

    if request.method == 'GET':
        # var_value = removeSpace("".join(var_value.splitlines()))
        formula = removeSpace("".join(formula.splitlines()))
        img_var = removeSpace("".join(img_var.splitlines()))

        if img_var != '---':
            var_value = var_value + ';' + img_var
        else:
            pass

        if arrange == "mapping":
            mapped_var = gen_mapped_var(var_value)

            # Check if item model contain additional formula
            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        elif arrange == "kombinasi":
            mapped_var = gen_combined_var(var_value)

            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        # format_img(mapped_var)
        pertanyaan = gen_question(stem, mapped_var)
        jawaban = gen_non_math_option(option, mapped_var)

        request.session['question'] = pertanyaan
        print(pertanyaan)
        request.session['answer'] = jawaban
        request.session['indikator_id'] = indikator_id
        request.session['inus_id'] = inus_id
        return render(request, 'resulttemplate.html')


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def SimpanTemplate(request):
    # stem = request.session.get('stem')
    # amount = request.session.get('amount')
    # arrange = request.session.get('arrange')
    # var_value = request.session.get('var_value')
    # var_name = request.session.get('var_name')
    # formula = request.session.get('formula')
    # option = request.session.get('option')

    # judul_template = request.GET.get('judultemplate')
    # p = TemplateSoal2(judul=judul_template, soal=stem, variasi=amount, pengacakan=arrange, var1=var_name, variabel=var_value,
    #                  rumus=formula, jawaban=option)
    # p.save()

    # return redirect('template')
    inus_id = request.session.get('inus_id')
    user_id = request.user.id
    id_user = User.objects.get(id=user_id)
    stem = request.session.get('stem')
    amount = request.session.get('amount')
    arrange = request.session.get('arrange')
    var_value = request.session.get('var_value')
    var_name = request.session.get('var_name')
    formula = request.session.get('formula')
    option = request.session.get('option')

    judul_template = request.GET.get('judultemplate')
    p = TemplateSoal(judul=judul_template, id_user=id_user, soal=stem, variasi=amount, pengacakan=arrange, var1=var_name, variabel=var_value,
                     rumus=formula, jawaban=option)
    p.save()

    return redirect('template')


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def Result(request):
    indikator_id = request.session.get('indikator_id')
    inus_id = request.session.get('inus_id')
    id_creator = request.user.id
    stem = request.session.get('stem')
    amount = request.session.get('amount')
    arrange = request.session.get('arrange')
    var_value = request.session.get('var_value')
    var_name = request.session.get('var_name')
    formula = request.session.get('formula')
    option = request.session.get('option')
    img_var = request.session.get('img_var')

    if request.method == 'GET':
        # var_value = removeSpace("".join(var_value.splitlines()))
        formula = removeSpace("".join(formula.splitlines()))
        img_var = removeSpace("".join(img_var.splitlines()))

        if img_var != '---':
            var_value = var_value + ';' + img_var
        else:
            pass

        if arrange == "mapping":
            mapped_var = gen_mapped_var(var_value)

            # Check if item model contain additional formula
            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        elif arrange == "kombinasi":
            mapped_var = gen_combined_var(var_value)

            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        # format_img(mapped_var)
        pertanyaan = gen_question(stem, mapped_var)
        jawaban = gen_non_math_option(option, mapped_var)

        request.session['question'] = pertanyaan
        request.session['answer'] = jawaban
        request.session['indikator_id'] = indikator_id
        request.session['inus_id'] = inus_id
        return render(request, 'result2.html')


@csrf_exempt
def Simpan(request):
    inus_id = request.session.get('inus_id')
    penugasan_id = request.session.get('penugasan_id')
    user_id = request.user.id
    id_creator = User.objects.get(id=user_id)
    stem = request.session.get('stem')
    amount = request.session.get('amount')
    arrange = request.session.get('arrange')
    var_value = request.session.get('var_value')
    var_name = request.session.get('var_name')
    formula = request.session.get('formula')
    option = request.session.get('option')
    img_var = request.session.get('img_var')
    status_soal = StatusSoal.objects.get(id=1)
    penugasan = Penugasan.objects.get(id=penugasan_id)
    if request.GET.get('type') == 'update':
        id_soal = request.GET.get('id')
        soal = Soal.objects.get(id=id_soal)
        soal.soal = stem
        soal.variasi = amount
        soal.pengacakan = arrange
        soal.var1 = var_name
        soal.variabel = var_value
        soal.rumus = formula
        soal.jawaban = option
        soal.variabel_gambar = img_var
        soal.indikator = penugasan  # edit later
        soal.created_date = datetime.now()
        soal.status = status_soal
        soal.save()

    else:
        nama_soal = request.GET.get('namasoal')
        p = Soal(judul=nama_soal, id_creator=id_creator, soal=stem, variasi=amount, pengacakan=arrange, var1=var_name, variabel=var_value,
                 rumus=formula, jawaban=option, variabel_gambar=img_var, indikator=penugasan, created_date=datetime.now(), status=status_soal)
        p.save()
        inus = Indikator_User.objects.get(id=inus_id)
        inus.is_done = 1
        inus.save()

    return redirect('pilihindikatorview')


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def DataSoal(request):

    if request.method == 'GET':
        id_creator = request.user.id
        penugasan_id = request.session.get('penugasan_id')
        get_penugasan_id = request.GET.get('penugasan_id')
        if get_penugasan_id is not None:
            penugasan_id = get_penugasan_id
        penugasan = Penugasan.objects.get(id=penugasan_id)
        soal = Soal.objects.filter(
            id_creator=id_creator, indikator=penugasan).order_by('id').reverse()
        id_user = ''
        pesan_revisi = ''

        for soals in soal:
            revisian = RevisiSoal.objects.filter(soal=soals)
            for revisi in revisian:
                id_user = revisi.id_user
                pesan_revisi = revisi.pesan
        context = {
            'id_user': id_user,
            'datasoal': soal,
            'soalsoal': 'active',
            'pesan_revisi': pesan_revisi,
            'revisian': revisian,
        }

        if request.user.is_authenticated:
            return render(request, 'datasoal.html', context)
        else:
            return redirect('index')

    elif request.method == 'POST':
        id_soal = request.POST.get('id')

        if request.POST.get('type') == 'lihat':
            soal = Soal.objects.get(id=id_soal)
            request.session['id_soal'] = id_soal
            request.session['stem'] = soal.soal
            request.session['amount'] = soal.variasi
            request.session['arrange'] = soal.pengacakan
            request.session['var_value'] = soal.variabel
            request.session['formula'] = soal.rumus
            request.session['var_name'] = soal.var1
            request.session['option'] = soal.jawaban
            request.session['title'] = soal.judul
            request.session['img_var'] = soal.variabel_gambar
            return redirect('hasil')

        elif request.POST.get('type') == 'hapus':
            inus_id = request.session.get('inus_id')
            inus = Indikator_User.objects.get(id=inus_id)
            inus.is_done = 1
            inus.save()
            inus.delete()
            Soal.objects.filter(id=id_soal).delete()
            return redirect('pilihindikatorview')

        elif request.POST.get('type') == 'edit':
            base_url = reverse('identitas')
            query_string = urlencode({
                'edit': 'true',
                'id': id_soal
            })
            url = '{}?{}'.format(base_url, query_string)

            soal = Soal.objects.get(id=id_soal)
            request.session['stem'] = soal.soal
            request.session['amount'] = soal.variasi
            request.session['arrange'] = soal.pengacakan
            request.session['var_value'] = soal.variabel
            request.session['formula'] = soal.rumus
            request.session['var_name'] = soal.var1
            request.session['option'] = soal.jawaban
            request.session['img_var'] = soal.variabel_gambar
            request.session['penugasan_id'] = soal.indikator_id
            return redirect(url)


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def Template(request):
    if request.method == 'GET':
        user_id = request.user.id
        soal = TemplateSoal.objects.filter(
            id_user_id=user_id).order_by('id').reverse()
        context = {
            'datasoal': soal,
            'template': 'active',
        }
        return render(request, 'template.html', context)

    elif request.method == 'POST':
        id_soal = request.POST.get('id')
        soal = TemplateSoal.objects.get(id=id_soal)
        request.session['stem'] = soal.soal
        request.session['amount'] = soal.variasi
        request.session['arrange'] = soal.pengacakan
        request.session['var_value'] = soal.variabel
        request.session['formula'] = soal.rumus
        request.session['var_name'] = soal.var1
        request.session['option'] = soal.jawaban
        request.session['title'] = soal.judul
        request.session['img_var'] = '---'

        if request.POST.get('type') == 'lihat':
            return redirect('hasil')
        elif request.POST.get('type') == 'gunakan':
            request.session['penugasan_id'] = None
            url = 'buatindikator'
            return redirect(url)


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def Hasil(request):
    id_creator = request.user.id
    id_soal = request.session.get('id_soal')
    stem = request.session.get('stem')
    amount = request.session.get('amount')
    arrange = request.session.get('arrange')
    var_value = request.session.get('var_value')
    var_name = request.session.get('var_name')
    formula = request.session.get('formula')
    option = request.session.get('option')
    img_var = request.session.get('img_var')
    # soal = Soal.objects.get(id=id_soal)
    # status = soal.status.id

    if request.method == 'GET':
        # var_value = removeSpace("".join(var_value.splitlines()))
        formula = removeSpace("".join(formula.splitlines()))
        if img_var:
            img_var = removeSpace("".join(img_var.splitlines()))

            if img_var != '---':
                var_value = var_value + ';' + img_var
            else:
                pass

        if arrange == "mapping":
            mapped_var = gen_mapped_var(var_value)

            # Check if item model contain additional formula
            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        elif arrange == "kombinasi":
            mapped_var = gen_combined_var(var_value)

            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        # format_img(mapped_var)
        pertanyaan = gen_question(stem, mapped_var)
        jawaban = gen_non_math_option(option, mapped_var)

        request.session['question'] = pertanyaan
        request.session['answer'] = jawaban
        return render(request, 'hasil.html')

    elif request.method == 'POST':

        if request.POST.get('type') == 'submit':
            status = StatusSoal.objects.get(id=2)
            soal = Soal.objects.get(id=id_soal)
            soal.status = status
            soal.save()
            # peninjaus = Peninjau.objects.filter(soal_id=soal.id)
            # for peninjau in peninjaus:
            #     peninjau.delete()

            for soal, jawaban in zip(request.session['question'], request.session['answer']):
                generatedSoal = GeneratedSoal()
                generatedSoal.generated_soal = soal
                generatedSoal.generated_option = jawaban
                generatedSoal.parent = Soal.objects.get(id=id_soal)
                generatedSoal.save()
            return redirect('datasoal')


@csrf_exempt
def Panduan(request):
    context = {
        'panduan': 'active',
    }
    if request.user.is_authenticated:
        return render(request, 'panduan.html', context)
    else:
        return redirect('index')


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def Hapus(request):
    id_soal = request.GET.get('id')
    Soal.objects.filter(id=id_soal).delete()

    return redirect('datasoal')


@csrf_exempt
def logoutView(request):
    logout(request)
    return redirect('index')


@csrf_exempt
def changePassword(request):

    if request.method == 'POST':
        form = PasswordChangeForm(request.POST)
        if form.is_valid():
            old_password = form.cleaned_data['oldpass']
            new_password = form.cleaned_data['newpass']
            confirm_password = form.cleaned_data['confirmpass']

            username = request.user.username
            user = authenticate(username=username, password=old_password)

            if user is not None:
                if new_password == confirm_password:
                    if new_password != old_password:
                        message = 'Change Password success'
                        u = User.objects.get(username=username)
                        u.set_password(new_password)
                        u.save()
                        update_session_auth_hash(request, u)
                    else:
                        message = 'Your New Password and Old Password are same!'
                else:
                    message = 'New Password and Confirm Password are different!'
            else:
                message = 'Your password is wrong'

            print(old_password, confirm_password, new_password)
            form = PasswordChangeForm()
            return render(request, 'changepassword.html', {'forms': form, 'err_message': message})

        print(old_password, new_password, confirm_password)

    else:
        form = PasswordChangeForm()

    return render(request, 'changepassword.html', {'forms': form})


@csrf_exempt
def Download(request):
    document = Document()
    for soal, t in zip(request.session['question'], request.session['answer']):
        paragraph = document.add_paragraph()
        if 'math-tex' in soal:
            cari = BeautifulSoup(soal).find("span", {"class": "math-tex"})
            delete = '<span class="math-tex">' + cari.text + '</span>'
            abcd = soal.replace(delete, ",|,")
            final = strip_tags(unescape(abcd)).split(",")
            a = cari.text.replace('\(', '')
            tex_eq = a.replace('\)', '')
            raw = str_to_raw(tex_eq)
            eq = mcv.convert(raw)
            tree = etree.fromstring(eq)
            xslt = etree.parse('MML2OMML.XSL')
            transform = etree.XSLT(xslt)
            new_dom = transform(tree)
            for x in range(0, len(final)):
                # if final[x] == "|":
                if 'codecogs' in str(cari[i]):
                    paragraph._element.append(new_dom.getroot())
                else:
                    paragraph.add_run(final[x].strip())
                    paragraph.add_run(str(cari[i]).strip())

        elif 'img' in soal:
            abcd = soal
            print(abcd)
            soup = BeautifulSoup(abcd, 'html.parser')
            cari = soup.find_all('img')
            print(cari)
            url = []
            response = []
            image = []
            image2 = []
            width_list = []
            height_list = []
            j = 0

            for i in range(soal.count('img')):
                if 'codecogs' in str(cari[i]):
                    abcd = soal
                    url.append(cari[i]['src'])
                    response.append(requests.get(url[i], stream=True))
                    print(url[i])
                    image.append(io.BytesIO(response[i].content))
                    image2.append(Image.open(image[i]))
                    width, height = image2[i].size
                    width_list.append(width)
                    height_list.append(height)
                else:
                    abcd = soal
                    url.append(cari[i]['src'])
                    response.append(requests.get(
                        'http://127.0.0.1:8000/' + url[i], stream=True))
                    print(url[i])
                    image.append(io.BytesIO(response[i].content))
                    image2.append(Image.open(image[i]))
                    width, height = image2[i].size
                    width_list.append(width)
                    height_list.append(height)

                if 'codecogs' or 'media' in str(cari[i]):
                    j += 1
                    if width_list[j-1] >= 288 and width_list[j-1] >= height_list[j-1]:
                        paragraph.add_run().add_picture(
                            image[j-1], width=Inches(3.0))
                        parsed = strip_tags(unescape(soal.strip()))
                        print(parsed + " ini adalah soal")
                        paragraph.add_run(parsed)
                    elif height_list[j-1] >= 192 and height_list[j-1] > width_list[j-1]:
                        paragraph.add_run().add_picture(
                            image[j-1], height=Inches(2.0))
                        parsed = strip_tags(unescape(soal.strip()))
                        paragraph.add_run(parsed)
                    else:
                        paragraph.add_run().add_picture(image[j-1])
                        parsed = strip_tags(unescape(soal.strip()))
                        paragraph.add_run(parsed)
                else:
                    # paragraph.add_run(final[x])
                    paragraph.add_run(str(cari[i]))

        else:
            parsed = strip_tags(unescape(soal.strip()))
            print(parsed + " ini adalah soal")
            paragraph.add_run(parsed)
        paragraph = document.add_paragraph()
        for jawaban in t:
            if 'img' in jawaban:
                abcd = jawaban
                cari = BeautifulSoup(abcd).find_all('img')
                url = []
                response = []
                image = []
                image2 = []
                width_list = []
                height_list = []
                j = 0

                for i in range(jawaban.count('img')):
                    if 'codecogs' in str(cari[i]):
                        abcd = jawaban
                        url.append(cari[i]['src'])
                        response.append(requests.get(url[i], stream=True))
                        print(url[i])
                        image.append(io.BytesIO(response[i].content))
                        image2.append(Image.open(image[i]))
                        width, height = image2[i].size
                        width_list.append(width)
                        height_list.append(height)
                    else:
                        abcd = soal
                        url.append(cari[i]['src'])
                        response.append(requests.get(
                            'http://127.0.0.1:8000/' + url[i], stream=True))
                        print(url[i])
                        image.append(io.BytesIO(response[i].content))
                        image2.append(Image.open(image[i]))
                        width, height = image2[i].size
                        width_list.append(width)
                        height_list.append(height)

                    if 'codecogs' or 'media' in str(cari[i]):
                        j += 1
                        if width_list[j-1] >= 288 and width_list[j-1] >= height_list[j-1]:
                            parsed = strip_tags(unescape(jawaban.strip()))
                            paragraph.add_run(parsed)
                            print(parsed + " ini adalah jawaban")
                            paragraph.add_run().add_picture(
                                image[j-1], width=Inches(3.0))
                        elif height_list[j-1] >= 192 and height_list[j-1] > width_list[j-1]:
                            paragraph.add_run().add_picture(
                                image[j-1], height=Inches(2.0))
                            parsed = strip_tags(unescape(soal.strip()))
                            paragraph.add_run(parsed)
                        else:
                            paragraph.add_run().add_picture(image[j-1])
                            parsed = strip_tags(unescape(soal.strip()))
                            paragraph.add_run(parsed)
                    else:
                        # paragraph.add_run(final[x])
                        paragraph.add_run(str(cari[i]))
            else:
                if '\(' in jawaban:
                    y = jawaban.replace('\(', '<span class="math-tex">')
                    z = y.replace('\)', '</span>')
                    cari2 = BeautifulSoup(z).find(
                        "span", {"class": "math-tex"})
                    delete = '<span class="math-tex">' + cari2.text + '</span>'
                    abcd = unescape(z).replace(delete, ",|,")
                    final = abcd.split(",")
                    tex_eq = cari2.text
                    raw = str_to_raw(tex_eq)
                    eq = mcv.convert(raw)
                    tree = etree.fromstring(eq)
                    xslt = etree.parse('MML2OMML.XSL')
                    transform = etree.XSLT(xslt)
                    new_dom = transform(tree)
                    for x in range(0, len(final)):
                        # if final[x] == "|":
                        if 'codecogs' in str(cari[i]):
                            paragraph._element.append(new_dom.getroot())
                        else:
                            # paragraph.add_run(final[x].strip())
                            paragraph.add_run(str(cari[i]).strip())
                    paragraph = document.add_paragraph()
                else:
                    parsed = strip_tags(unescape(jawaban.strip()))
                    paragraph.add_run(parsed)
                    paragraph = document.add_paragraph()

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename= ' + \
        request.session['title'] + '.docx'
    file_stream = io.BytesIO()
    document.save(response)
    document.save(file_stream)

    return response


@csrf_exempt
def Kirim(request):
    document = Document()

    api_upload = "https://moodle.aigyujiem.com/webservice/restful/server.php/core_files_upload"
    api_copy_private = "https://moodle.aigyujiem.com/webservice/restful/server.php/core_user_add_user_private_files"

    for soal, t in zip(request.session['question'], request.session['answer']):
        paragraph = document.add_paragraph()
        if 'math-tex' in soal:
            cari = BeautifulSoup(soal).find("span", {"class": "math-tex"})
            delete = '<span class="math-tex">' + cari.text + '</span>'
            abcd = soal.replace(delete, ",|,")
            final = strip_tags(unescape(abcd)).split(",")
            a = cari.text.replace('\(', '')
            tex_eq = a.replace('\)', '')
            raw = str_to_raw(tex_eq)
            eq = mcv.convert(raw)
            tree = etree.fromstring(eq)
            xslt = etree.parse('MML2OMML.XSL')
            transform = etree.XSLT(xslt)
            new_dom = transform(tree)
            for x in range(0, len(final)):
                if final[x] == "|":
                    paragraph._element.append(new_dom.getroot())
                else:
                    paragraph.add_run(final[x].strip())

        elif 'img' in soal:
            abcd = soal
            cari = BeautifulSoup(abcd).find_all('img')
            url = []
            response = []
            image = []
            image2 = []
            width_list = []
            height_list = []
            j = 0

            for i in range(soal.count('img')):
                if 'codecogs' in str(cari[i]):
                    delete = str(cari[i]).replace('/>', ' />')
                else:
                    delete = str(cari[i]).replace('/>', '>')
                abcd = abcd.replace(delete, ",|,")
                final = strip_tags(unescape(abcd)).split(",")
                url.append(cari[i]['src'])
                response.append(requests.get(url[i], stream=True))
                image.append(io.BytesIO(response[i].content))
                image2.append(Image.open(image[i]))
                width, height = image2[i].size
                width_list.append(width)
                height_list.append(height)
            for x in range(0, len(final)):
                if final[x] == "|":
                    j += 1
                    if width_list[j-1] >= 288 and width_list[j-1] >= height_list[j-1]:
                        paragraph.add_run().add_picture(
                            image[j-1], width=Inches(3.0))
                    elif height_list[j-1] >= 192 and height_list[j-1] > width_list[j-1]:
                        paragraph.add_run().add_picture(
                            image[j-1], height=Inches(2.0))
                    else:
                        paragraph.add_run().add_picture(image[j-1])
                else:
                    paragraph.add_run(final[x])

        else:
            parsed = strip_tags(unescape(soal.strip()))
            paragraph.add_run(parsed)
        paragraph = document.add_paragraph()
        for jawaban in t:
            if 'img' in jawaban:
                abcd = jawaban
                cari = BeautifulSoup(abcd).find_all('img')
                url = []
                response = []
                image = []
                image2 = []
                width_list = []
                height_list = []
                j = 0

                for i in range(jawaban.count('img')):
                    if 'codecogs' in str(cari[i]):
                        delete = str(cari[i]).replace('/>', ' />')
                    else:
                        delete = str(cari[i]).replace('/>', '>')
                    abcd = abcd.replace(delete, ",|,")
                    final = strip_tags(unescape(abcd)).split(",")
                    url.append(cari[i]['src'])
                    response.append(requests.get(url[i], stream=True))
                    image.append(io.BytesIO(response[i].content))
                    image2.append(Image.open(image[i]))
                    width, height = image2[i].size
                    width_list.append(width)
                    height_list.append(height)
                for x in range(0, len(final)):
                    if final[x] == "|":
                        j += 1
                        if width_list[j-1] >= 288 and width_list[j-1] >= height_list[j-1]:
                            paragraph.add_run().add_picture(
                                image[j-1], width=Inches(3.0))
                        elif height_list[j-1] >= 192 and height_list[j-1] > width_list[j-1]:
                            paragraph.add_run().add_picture(
                                image[j-1], height=Inches(2.0))
                        else:
                            paragraph.add_run().add_picture(image[j-1])
                    else:
                        paragraph.add_run(final[x])
            else:
                if '\(' in jawaban:
                    y = jawaban.replace('\(', '<span class="math-tex">')
                    z = y.replace('\)', '</span>')
                    cari2 = BeautifulSoup(z).find(
                        "span", {"class": "math-tex"})
                    delete = '<span class="math-tex">' + cari2.text + '</span>'
                    abcd = unescape(z).replace(delete, ",|,")
                    final = abcd.split(",")
                    tex_eq = cari2.text
                    raw = str_to_raw(tex_eq)
                    eq = mcv.convert(raw)
                    tree = etree.fromstring(eq)
                    xslt = etree.parse('MML2OMML.XSL')
                    transform = etree.XSLT(xslt)
                    new_dom = transform(tree)
                    for x in range(0, len(final)):
                        if final[x] == "|":
                            paragraph._element.append(new_dom.getroot())
                        else:
                            paragraph.add_run(final[x].strip())
                    paragraph = document.add_paragraph()
                else:
                    parsed = strip_tags(unescape(jawaban.strip()))
                    paragraph.add_run(parsed)
                    paragraph = document.add_paragraph()

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename= ' + \
        request.session['title'] + '.docx'
    file_stream = io.BytesIO()
    document.save(response)
    document.save(file_stream)
    file_bytes = file_stream.getvalue()
    base64_encoded_data = base64.b64encode(file_bytes)
    base64_message = base64_encoded_data.decode(('utf-8'))

    header = {
        'Authorization': request.session['moodle_token'],
        'Accept': 'application/json',
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.97 Safari/537.36"
    }

    payload_send = {
        "component": "user",
        "filearea": "draft",
        "itemid": 0,
        "filepath": "/",
        "filename": request.session['title'] + '.docx',
        "contextlevel": "user",
        "instanceid": request.session['user_id'],
        "filecontent": base64_message
    }

    moodle_response_send = requests.post(
        api_upload, headers=header, data=payload_send)
    request.session['file_id'] = moodle_response_send.json()['itemid']

    payload_copy = {
        "draftid": request.session['file_id']
    }

    moodle_response_copy = requests.post(
        api_copy_private, headers=header, data=payload_copy)

    return redirect('hasilpaket')


@csrf_exempt
def str_to_raw(s):
    raw_map = {8: r'\b', 7: r'\a', 12: r'\f',
               10: r'\n', 13: r'\r', 9: r'\t', 11: r'\v'}
    return r''.join(i if ord(i) > 32 else raw_map.get(ord(i), i) for i in s)


# PJM Stuffs
@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def indikator_view(request):
    if request.method == 'GET':
        id_creator = request.user.id
        penugasan_query = Penugasan.objects.filter(
            creator=id_creator).order_by('id').reverse()
        penugasan_data = []
        order = 1
        for tugas in penugasan_query:

            # All field needed
            id = tugas.id
            kurikulum = tugas.kurikulum
            matkul = tugas.mata_kuliah
            so = tugas.student_outcome
            lo = tugas.learning_outcome
            topik = tugas.topik
            soal = Soal.objects.filter(indikator=id).count()
            detail = {
                'order': order,
                'id': id,
                'kurikulum': kurikulum,
                'mata_kuliah': matkul,
                'student_outcome': so,
                'learning_outcome': lo,
                'topik': topik,
                'count': soal
            }

            penugasan_data.append(detail)
            order -= - 1
        data = {
            'tugas_data': penugasan_data
        }
        return render(request, 'indikator.html', data)

    elif request.method == 'POST':

        id_penugasan = request.POST.get('id')

        if request.POST.get('type') == 'detail':
            request.session['id'] = id_penugasan
            return redirect('detailindikator')


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def buatIndikator(request):
    form = IndikatorForm()
    user = User.objects.get(id=request.user.id)
    if request.method == 'POST':
        form = IndikatorForm(request.POST)
        if form.is_valid():
            obj_indikator = form.save(commit=False)
            obj_indikator.creator = request.user
            obj_indikator.created_date = datetime.now()
            obj_indikator.save()
            penugasan = Penugasan.objects.get(id=obj_indikator.id)
            obj_inus = Indikator_User()
            obj_inus.penugasan = penugasan
            obj_inus.learning_outcome = penugasan.learning_outcome
            obj_inus.user = user
            obj_inus.save()
            inus_id = obj_inus.id
            penugasan_id = obj_indikator.id
            request.session['inus_id'] = inus_id
            if request.session.get('penugasan_id') == None:
                request.session['penugasan_id'] = penugasan_id
                base_url = reverse('identitas')
                query_string = urlencode({
                    'fromtemplate': 'true'
                })
                url = '{}?{}'.format(base_url, query_string)
                return redirect(url)
            request.session['penugasan_id'] = penugasan_id

            return redirect('identitas')
    else:
        form = IndikatorForm()
    return render(request, 'buatindikator.html', {'forms': form, 'buatindikator': 'active'})


@csrf_exempt
def StudentOutcome_TopicUpdate(request, pk):
    penugasan = get_object_or_404(Penugasan, pk=pk)
    if request.method == 'POST':
        form = IndikatorForm(request.POST, instance=penugasan)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.creator = request.user
            obj.created_date = datetime.now
            obj.save()
            return redirect('SO_TP_change', pk=pk)
    return render(request, 'buatindikator.html', {'forms': form})


@csrf_exempt
def LearningOutcome_Update(request, pk):
    penugasan = get_object_or_404(Penugasan, pk=pk)
    if request.method == 'POST':
        form = IndikatorForm(request.POST, instance=penugasan)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.creator = request.user
            obj.created_date = datetime.now
            obj.save()
            return redirect('LO_change', pk=pk)
    return render(request, 'buatindikator.html', {'forms': form})


@csrf_exempt
def load_LO(request):
    so_id = request.GET.get('studentOutcome_id')
    learningOutcome = LearningOutcome.objects.filter(SO_id=so_id).all()
    return render(request, 'lo_dropdown_list_update.html', {'Lo': learningOutcome})


@csrf_exempt
def load_SO(request):
    mk_id = request.GET.get('matkul_id')
    studentOutcome = StudentOutcome.objects.filter(matkul_id=mk_id).all()
    topik = Topic.objects.filter(matkul_id=mk_id).all()
    return render(request, 'so_dropdown_list_update.html', {'So': studentOutcome})


@csrf_exempt
def load_TP(request):
    mk_id = request.GET.get('matkul_id')
    topik = Topic.objects.filter(matkul_id=mk_id).all()
    print(mk_id)
    return render(request, 'tp_dropdown_list_update.html', {'TP': topik})


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def detailIndikator(request):
    penugasan_id = request.session.get('id')
    penugasan = Penugasan.objects.get(id=penugasan_id)

    soal_query = Soal.objects.filter(
        indikator=penugasan_id).order_by('id').reverse()
    soal_data = []

    already_assign_query = Indikator_User.objects.filter(penugasan=penugasan)
    already_assign_data = []

    learning_outcome = penugasan.learning_outcome
    student_outcome = penugasan.student_outcome

    id = penugasan.id
    matakuliah = penugasan.mata_kuliah
    materi = penugasan.topik
    kurikulum = penugasan.kurikulum

    indikator_data = [{
        'id': id,
        'learning_outcome': learning_outcome,
        'student_outcome': student_outcome,
        'topik': materi,
        'mata_kuliah': matakuliah,
        'kurikulum': kurikulum,
    }]

    flag = 0
    for soal in soal_query:
        id = soal.id
        status = soal.status
        creator = soal.id_creator
        created_date = soal.created_date
        judul = soal.judul

        detail = {
            'id': id,
            'status': status,
            'creator': creator,
            'created_date': created_date,
            'judul': judul,
        }

        soal_data.append(detail)
        flag = 1

    for already_assign in already_assign_query:
        front_name = already_assign.user.first_name
        last_name = already_assign.user.last_name
        already_assign_data.append(front_name + ' ' + last_name)

    if request.method == 'POST':
        if request.POST.get('type') == 'assignUser':
            request.session['id'] = penugasan_id
            form = AssignUser(request.POST)
            if form.is_valid():
                user = form.cleaned_data['user']
                penugasan = Penugasan.objects.get(id=penugasan_id)
                already_assign = Indikator_User.objects.filter(
                    user=user, penugasan=penugasan)
                if len(already_assign) == 0:
                    obj = Indikator_User()
                    obj.penugasan = penugasan
                    obj.learning_outcome = penugasan.learning_outcome
                    obj.user = user
                    obj.save()
                else:
                    pass

            return redirect('detailindikator')

        elif request.POST.get('type') == 'review':
            id_soal = request.POST.get('id_soal')
            soal = Soal.objects.get(id=id_soal)
            request.session['id'] = penugasan_id
            request.session['id_soal'] = id_soal
            request.session['stem'] = soal.soal
            request.session['amount'] = soal.variasi
            request.session['arrange'] = soal.pengacakan
            request.session['var_value'] = soal.variabel
            request.session['formula'] = soal.rumus
            request.session['var_name'] = soal.var1
            request.session['option'] = soal.jawaban
            request.session['title'] = soal.judul
            request.session['img_var'] = soal.variabel_gambar
            return redirect('reviewindikator')
    else:
        form = AssignUser()

    data = {
        'indikator_data': indikator_data,
        'soal_data': soal_data,
        'flag': flag,
        'form': form,
        'already_assign': already_assign_data,
    }

    return render(request, 'detailindikator.html', data)


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def reviewIndikator(request):
    id_creator = request.user.id
    id_soal = request.session.get('id_soal')
    stem = request.session.get('stem')
    amount = request.session.get('amount')
    arrange = request.session.get('arrange')
    var_value = request.session.get('var_value')
    var_name = request.session.get('var_name')
    formula = request.session.get('formula')
    option = request.session.get('option')
    img_var = request.session.get('img_var')
    soal = Soal.objects.get(id=id_soal)
    status = soal.status.id

    if request.method == 'GET':
        # var_value = removeSpace("".join(var_value.splitlines()))
        formula = removeSpace("".join(formula.splitlines()))
        img_var = removeSpace("".join(img_var.splitlines()))

        if img_var != '---':
            var_value = var_value + ';' + img_var
        else:
            pass

        if arrange == "mapping":
            mapped_var = gen_mapped_var(var_value)

            # Check if item model contain additional formula
            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        elif arrange == "kombinasi":
            mapped_var = gen_combined_var(var_value)

            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        # format_img(mapped_var)
        pertanyaan = gen_question(stem, mapped_var)
        jawaban = gen_non_math_option(option, mapped_var)

        form = RevisiSoalForm()
        request.session['question'] = pertanyaan
        request.session['answer'] = jawaban
        return render(request, 'reviewindikator.html', {'status': status, 'form': form})

    elif request.method == 'POST':
        if request.POST.get('type') == 'terima':

            status = StatusSoal.objects.get(id=3)
            soal = Soal.objects.get(id=id_soal)
            soal.status = status
            soal.save()

            for soal, jawaban in zip(request.session['question'], request.session['answer']):
                generatedSoal = GeneratedSoal()
                generatedSoal.generated_soal = soal
                generatedSoal.generated_option = jawaban
                generatedSoal.parent = Soal.objects.get(id=id_soal)
                generatedSoal.save()

            return redirect('detailindikator')

        elif request.POST.get('type') == 'tolak':
            form = RevisiSoalForm(request.POST)
            user_id = request.user.id
            if form.is_valid():
                pesan_revisi = form.cleaned_data['pesan_revisi']
                soal = Soal.objects.get(id=id_soal)
                obj = RevisiSoal()
                obj.id_user = User.objects.get(id=user_id)
                obj.soal = soal
                obj.pesan = pesan_revisi
                obj.save()
                status = StatusSoal.objects.get(id=2)
                soal.status = status
                soal.save()
                return redirect('detailindikator')

        elif request.POST.get('type') == 'download':
            parent = Soal.objects.get(id=id_soal)
            generated_soal = GeneratedSoal.objects.filter(parent=parent)
            question_data = []
            option_data = []
            judul = parent.judul

            for soal in generated_soal:
                question = soal.generated_soal
                jawaban = literal_eval(soal.generated_option)
                question_data.append(question)
                option_data.append(jawaban)

            request.session['question'] = question_data
            request.session['answer'] = option_data
            request.session['title'] = judul

            return redirect('test')


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def generateDetail(request):
    return render(request, 'generatepaketdetail.html')


@csrf_exempt
def pilihindikator(request):
    if request.method == 'GET':
        user_id = request.user.id
        penugasan_user_query = Indikator_User.objects.filter(
            user=user_id, is_done=0).order_by('id')
        penugasan_data = []
        penugasan = penugasan_user_query

        for penugasan_instance in penugasan_user_query:
            penugasan = penugasan_instance.penugasan
            instance_id = penugasan_instance.id
            learning_outcome = penugasan.learning_outcome
            student_outcome = penugasan.student_outcome

            id = penugasan.id
            topik = penugasan.topik
            kurikulum = penugasan.kurikulum
            matakuliah = penugasan.mata_kuliah

            detail = {
                'inus_id': instance_id,
                'id': id,
                'mata_kuliah': matakuliah,
                'student_outcome': student_outcome,
                'learning_outcome': learning_outcome,
                'topik': topik,
                'kurikulum': kurikulum,
            }
            penugasan_data.append(detail)

    elif request.method == 'POST':
        if request.POST.get('type') == 'kerjakan':
            inus_id = request.POST.get('inus_id')
            penugasan_id = request.POST.get('id')
            request.session['inus_id'] = inus_id
            request.session['penugasan_id'] = penugasan_id
            return redirect('identitas')

    return render(request, 'pilihindikator.html', {'penugasan_data': penugasan_data})


@csrf_exempt
def pilihIndikatorView(request):
    if request.method == 'GET':
        user_id = request.user.id
        penugasan_user_query = Indikator_User.objects.filter(
            user=user_id, is_done=1).order_by('id')  # copy this line for peer review
        penugasan_data = []
        penugasan = penugasan_user_query

        for penugasan_instance in penugasan_user_query:
            penugasan = penugasan_instance.penugasan

            instance_id = penugasan_instance.id
            student_outcome = penugasan.student_outcome
            learning_outcome = penugasan.learning_outcome

            id = penugasan.id
            topik = penugasan.topik
            kurikulum = penugasan.kurikulum
            matakuliah = penugasan.mata_kuliah

            detail = {
                'inus_id': instance_id,
                'id': id,
                'mata_kuliah': matakuliah,
                'student_outcome': student_outcome,
                'learning_outcome': learning_outcome,
                'topik': topik,
                'kurikulum': kurikulum,
            }
            penugasan_data.append(detail)

    elif request.method == 'POST':
        if request.POST.get('type') == 'lihat':
            inus_id = request.POST.get('inus_id')
            penugasan_id = request.POST.get('id')
            request.session['inus_id'] = inus_id
            request.session['penugasan_id'] = penugasan_id
            return redirect('datasoal')

    return render(request, 'pilihindikatorview.html', {'penugasan_data': penugasan_data, 'pilihindikatorview': 'active'})


@csrf_exempt
def buatPaketSoal(request):
    if request.method == 'POST':

        user = request.user
        form = IdentitasPaketSoalForm(request.POST)

        if form.is_valid():
            judul = form.cleaned_data['judul']
            kurikulum = form.cleaned_data['kurikulum']
            jumlah_paket = form.cleaned_data['jumlah_paket']
            tahun_ajaran = form.cleaned_data['tahun_ajaran']
            semester = form.cleaned_data['semester']

            obj = PaketSoal()
            obj.judul = judul
            obj.kurikulum = kurikulum
            obj.jumlah_paket = jumlah_paket
            obj.created_date = datetime.now()
            obj.creator = user
            obj.is_generated = 0
            obj.tahun_ajaran = tahun_ajaran
            obj.semester = semester
            obj.status = StatusPaketSoal.objects.get(id=1)
            obj.save()

            paket_soal_id = obj.id
            request.session['paket_soal_id'] = paket_soal_id
            return redirect('generatepaket')

    # if a GET (or any other method) we'll create a blank form
    else:
        form = IdentitasPaketSoalForm()

    return render(request, 'buatpaketsoal.html', {'forms': form})


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def DataPaket(request):
    if request.method == 'GET':
        user = request.user
        paket_soal_query = PaketSoal.objects.filter(
            creator=user).order_by('id')
        paket_soal_data = []

        for paket_soal in paket_soal_query:
            judul = paket_soal.judul

            created_date = paket_soal.created_date
            is_generated = paket_soal.is_generated
            generated_date = paket_soal.generated_date
            kurikulum = paket_soal.kurikulum

            id = paket_soal.id

            detail = {
                'judul': judul,
                'id': id,
                'created_date': created_date,
                'kurikulum': kurikulum,
                'is_generated': is_generated,
                'generated_date': generated_date
            }

            paket_soal_data.append(detail)

    elif request.method == 'POST':
        paket_soal_id = request.POST.get('id')

        if request.POST.get('type') == 'detail':
            request.session['paket_soal_id'] = paket_soal_id

            paket_soal = PaketSoal.objects.get(id=paket_soal_id)
            if paket_soal.is_generated == 0:
                return redirect('generatepaket')
            elif paket_soal.is_generated == 1:
                return redirect('hasilpaket')

        elif request.POST.get('type') == 'hapus':
            PaketSoal.objects.filter(id=paket_soal_id).delete()
            return redirect('datapaket')

    return render(request, 'datapaket.html', {'paket_soal': paket_soal_data, 'datapaket': 'active'})


def lihatReview(request):
    if request.method == "GET":
        paket_soal_id = request.session.get('paket_soal_id')
        paket_soal = PaketSoal.objects.get(id=paket_soal_id)
        brother_id = request.session.get('brother_id')
        print(brother_id)
        brother_soal = Soal_PaketSoal.objects.get(id=brother_id)
        reviewers = Peninjau.objects.filter(paket_soal=paket_soal)
        hasil_reviews = []
        for reviewer in reviewers:
            hasil_review = HasilReview.objects.get(
                soal_paketSoal=brother_soal, peninjau=reviewer)
            hasil_reviews.append(hasil_review)

    return render(request, 'lihatreview.html', {'hasil_reviews': hasil_reviews})


@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def generatePaket(request):
    paket_soal_id = request.session.get('paket_soal_id')
    paket_soal = PaketSoal.objects.get(id=paket_soal_id)
    daftar_soal = []

    brother_soal_query = Soal_PaketSoal.objects.filter(
        paket_soal=paket_soal, nomer_paket=1)
    jumlah_parent = Soal_PaketSoal.objects.filter(
        paket_soal=paket_soal, nomer_paket=1).count()

    for brother_soal in brother_soal_query:

        parentsoal = brother_soal.generated_soal.parent.soal
        indikatorsoal = brother_soal.generated_soal.parent.indikator
        id = brother_soal.id
        hasil_reviews = brother_soal.status

        detail = {
            'id': id,
            'indikator': indikatorsoal,
            'parentsoal': parentsoal,
            'review': hasil_reviews
        }

        daftar_soal.append(detail)

    if request.method == 'POST':
        if request.POST.get('type') == 'generate':
            paket_soal_id = request.session.get('paket_soal_id')
            paket_soal = PaketSoal.objects.get(id=paket_soal_id)
            jumlah_paket_soal = paket_soal.jumlah_paket
            daftar_soal = []

            brother_soal_query = Soal_PaketSoal.objects.filter(
                paket_soal=paket_soal, nomer_paket=1)
            reviewers = Peninjau.objects.filter(paket_soal=paket_soal)

            for brother_soal in brother_soal_query:
                for reviewer in reviewers:
                    hasil_reviews = HasilReview.objects.filter(
                        peninjau=reviewer, soal_paketSoal=brother_soal)
                    hasil_reviews.delete()
                nomer_urut = brother_soal.nomer_urut
                soal_parent_query = Soal_PaketSoal.objects.filter(
                    paket_soal=paket_soal, nomer_urut=nomer_urut)
                jumlah_soal_parent = Soal_PaketSoal.objects.filter(
                    paket_soal=paket_soal, nomer_urut=nomer_urut).count()
                soal_data = []

                for soal in soal_parent_query:
                    generated_soal_id = soal.generated_soal.id
                    soal_data.append(generated_soal_id)

                while jumlah_soal_parent < jumlah_paket_soal:
                    soal_id = random.choice(soal_data)
                    soal = GeneratedSoal.objects.get(id=soal_id)

                    nomer_paket = jumlah_soal_parent+1
                    obj = Soal_PaketSoal()
                    obj.generated_soal = soal
                    obj.paket_soal = paket_soal
                    obj.nomer_urut = nomer_urut
                    obj.nomer_paket = nomer_paket
                    obj.save()

                    jumlah_soal_parent = nomer_paket

                urutan = list(range(1, jumlah_paket_soal+1))

                soal_parent_query = Soal_PaketSoal.objects.filter(
                    paket_soal=paket_soal, nomer_urut=nomer_urut)
                for soal in soal_parent_query:
                    nomer_paket = random.choice(urutan)
                    soal.nomer_paket = nomer_paket
                    soal.save()
                    urutan.remove(nomer_paket)

            paket_soal.is_generated = 1
            paket_soal.generated_date = datetime.now()
            paket_soal.save()

            reviewers.delete()

            return redirect('hasilpaket')

        elif request.POST.get('type') == 'hapus_parent':
            paket_soal_id = paket_soal_id = request.session.get(
                'paket_soal_id')
            paket_soal = PaketSoal.objects.get(id=paket_soal_id)
            brother_id = request.POST.get('id')
            soal_brother = Soal_PaketSoal.objects.get(id=brother_id)
            nomer_urut = soal_brother.nomer_urut
            soal_query = Soal_PaketSoal.objects.filter(
                paket_soal=paket_soal, nomer_urut=nomer_urut).delete()

            return redirect('generatepaket')

        elif request.POST.get('type') == 'hapus':
            paket_soal_id = request.session.get('paket_soal_id')
            PaketSoal.objects.filter(id=paket_soal_id).delete()
            return redirect('datapaket')

        elif request.POST.get('type') == 'detail':
            brother_id = request.POST.get('id')
            request.session['brother_id'] = brother_id
            return redirect('soaldipilih')

        elif request.POST.get('type') == 'lihat_review':
            brother_id = request.POST.get('id')
            request.session['brother_id'] = brother_id

            return redirect('lihat_review')

        elif request.POST.get('type') == 'review':
            paket_soal_id = request.session.get('paket_soal_id')
            paket_soal = PaketSoal.objects.get(id=paket_soal_id)
            paket_soal.status = StatusPaketSoal.objects.get(id=2)
            paket_soal.save()

    return render(request, 'generatepaket.html', {'paket_soal': paket_soal, 'data_soal': daftar_soal, 'jumlah_soal': jumlah_parent})


@csrf_exempt
def soalDipilih(request):
    if request.method == "GET":
        brother_id = request.session.get('brother_id')
        paket_soal_id = request.session.get('paket_soal_id')
        soal_brother = Soal_PaketSoal.objects.get(id=brother_id)
        paket_soal = PaketSoal.objects.get(id=paket_soal_id)
        nomer_urut = soal_brother.nomer_urut

        soal_dipilih_query = Soal_PaketSoal.objects.filter(
            paket_soal=paket_soal, nomer_urut=nomer_urut)
        soal_dipilih_data = []

        for soal_dipilih in soal_dipilih_query:
            question = soal_dipilih.generated_soal.generated_soal
            option = soal_dipilih.generated_soal.generated_option

            detail = {
                'question': question,
                'option': literal_eval(option)
            }

            soal_dipilih_data.append(detail)

    return render(request, 'gp_soaldipilih.html', {'soal_data': soal_dipilih_data})


@csrf_exempt
def tambahSoal(request):
    if request.method == 'GET':
        creator = request.user
        soal_query = Soal.objects.filter(id_creator=creator.id, status=2)
        penugasan_data = []
        order = 1

        for soal in soal_query:

            penugasan_id = str(soal.indikator)
            penugasan_query = Penugasan.objects.filter(
                creator=creator, id=soal.indikator.id).order_by('id').reverse()

            for penugasan in penugasan_query:
                # Foreign
                student_outcome = penugasan.student_outcome
                learning_outcome = penugasan.learning_outcome

            # All field needed
                id = penugasan.id
                topik = penugasan.topik
                kurikulum = penugasan.kurikulum
                jumlah_soal = Soal.objects.filter(indikator=id).count()
                matkul = penugasan.mata_kuliah
                detail = {
                    'order': order,
                    'id': id,
                    'penugasan': penugasan,
                    'student_outcome': student_outcome,
                    'learning_outcome': learning_outcome,
                    'topik': topik,
                    'kurikulum': kurikulum,
                    'matakuliah': matkul,
                    'count': jumlah_soal,
                    'penulis': soal.id_creator
                }
                penugasan_data.append(detail)
                order -= - 1

        data = {
            'penugasan_data': penugasan_data,
        }

        return render(request, 'gp_tambahsoal.html', data)

    elif request.method == 'POST':

        if request.POST.get('type') == 'pilih':
            penugasan_id = request.POST.get('id')
            request.session['penugasan'] = penugasan_id
            return redirect('pilihparentsoal')


@csrf_exempt
def pilihParentSoal(request):
    if request.method == 'POST':
        if request.POST.get('type') == 'pilih':
            parent_id = request.POST.get('id')
            request.session['parent'] = parent_id
            request.session['overload'] = ''
            return redirect('pilihsoal')

    elif request.method == 'GET':
        penugasan_id = request.session['penugasan']
        penugasan = Penugasan.objects.get(id=penugasan_id)
        parent_soal_query = Soal.objects.filter(
            indikator=penugasan, status=2).order_by('id').reverse()
        parent_soal_data = []

        for parent_soal in parent_soal_query:
            id = parent_soal.id
            creator = parent_soal.id_creator
            created_date = parent_soal.created_date
            judul = parent_soal.judul
            stem = parent_soal.soal

            detail = {
                'id': id,
                'creator': creator,
                'created_date': created_date,
                'judul': judul,
                'stem': stem,
            }

            parent_soal_data.append(detail)

    return render(request, 'gp_pilihparentsoal.html', {'parent_soal_data': parent_soal_data})


@csrf_exempt
def pilihSoal(request):
    if request.method == 'POST':
        paket_soal_id = request.session['paket_soal_id']
        paket_soal = PaketSoal.objects.get(id=paket_soal_id)

        jumlah_paket = paket_soal.jumlah_paket

        soal_dipilih = request.POST.getlist('selected')

        if len(soal_dipilih) > jumlah_paket:
            request.session['overload'] = 'Soal yang dipilih melebihi jumlah paket'
            return redirect('pilihsoal')

        else:
            nomer_urut = (Soal_PaketSoal.objects.filter(
                paket_soal=paket_soal, nomer_paket=1).count())+1
            nomer_paket = 1
            for soal_id in soal_dipilih:
                soal = GeneratedSoal.objects.get(id=soal_id)

                obj = Soal_PaketSoal()
                obj.generated_soal = soal
                obj.paket_soal = paket_soal
                obj.nomer_urut = nomer_urut
                obj.nomer_paket = nomer_paket
                obj.save()

                if (paket_soal.status.id != 1 and nomer_paket == 1):
                    reviewers = Peninjau.objects.filter(paket_soal=paket_soal)
                    for reviewer in reviewers:
                        hasil_review = HasilReview()
                        hasil_review.status = StatusPaketSoal.objects.get(id=1)
                        hasil_review.peninjau = reviewer
                        hasil_review.soal_paketSoal = obj
                        hasil_review.save()
                        profile_reviewer = Profile.objects.get(
                            user=reviewer.peninjau_user)
                        if profile_reviewer.subskripsi_email:
                            threading.Thread(target=send_reviewer_notif_new, args=(
                                reviewer.peninjau_user.email,
                                f"{reviewer.peninjau_user.first_name} {reviewer.peninjau_user.last_name}",
                                f"{request.user.first_name} {request.user.last_name}",
                                paket_soal.judul,
                                nomer_urut,
                                request.build_absolute_uri('/soal-peninjau/'),
                                request.build_absolute_uri('/subscription/')
                            )).start()
                    paket_soal.status = StatusPaketSoal.objects.get(id=2)
                    paket_soal.save()
                nomer_paket += 1

            return redirect('generatepaket')

    elif request.method == 'GET':
        parent_id = request.session['parent']
        parent = Soal.objects.get(id=parent_id)
        error = request.session['overload']
        soal_query = GeneratedSoal.objects.filter(
            parent=parent).order_by('id').reverse()
        soal_data = []

        for soal in soal_query:
            id = soal.id
            question = soal.generated_soal
            option = literal_eval(soal.generated_option)

            detail = {
                'id': id,
                'question': question,
                'option': option,
            }

            soal_data.append(detail)

    return render(request, 'gp_pilihsoal.html', {'error': error, 'soal_data': soal_data})


@csrf_exempt
def PeerReview(request):
    if request.method == 'GET':
        user_id = request.user.id
        penugasan_query = Penugasan.objects.all()
        # print(penugasan_query)
        penugasan_data = []
        penugasan = ''
        i = 0

        for tugas in penugasan_query:
            penugasan_id = tugas.id
            penugasan_user_query = Indikator_User.objects.filter(
                is_done=1, penugasan=penugasan_id).exclude(user=user_id).order_by('id').reverse()
            if penugasan_user_query:

                for penugasan_instance in penugasan_user_query:
                    penugasan = penugasan_instance.penugasan
                    instance_id = penugasan_instance.id
                    student_outcome = penugasan.student_outcome
                    learning_outcome = penugasan.learning_outcome

                    id = penugasan.id
                    topik = penugasan.topik
                    kurikulum = penugasan.kurikulum
                    matakuliah = penugasan.mata_kuliah

                    detail = {
                        'inus_id': instance_id,
                        'id': id,
                        'mata_kuliah': matakuliah,
                        'student_outcome': student_outcome,
                        'learning_outcome': learning_outcome,
                        'topik': topik,
                        'kurikulum': kurikulum,
                    }
                if not penugasan_data:
                    penugasan_data.append(detail)
                elif id != penugasan_data[i]["id"]:
                    penugasan_data.append(detail)
                    i += 1

    elif request.method == 'POST':
        if request.POST.get('type') == 'lihat':
            penugasan_id = request.POST.get('id')
            request.session['penugasan_id'] = penugasan_id
            return redirect('list-soal')

    return render(request, 'peerreview.html', {'penugasan_data': penugasan_data, 'peerreview': 'active'})


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def PeerSoal(request):
    form = RevisiSoalForm()
    soal_id = request.session.get('soal_id')
    soal = Soal.objects.filter(id=soal_id)
    id_user = ''
    pesan_revisi = ''
    revisibaru = None
    revisian = ""
    if request.method == 'GET':
        id_soal = request.GET.get('soal_id')
        if id_soal is not None:
            soal = Soal.objects.filter(id=id_soal)
        for soals in soal:
            revisian = RevisiSoal.objects.filter(soal=soals)
            for revisi in revisian:
                id_user = revisi.id_user
                pesan_revisi = revisi.pesan
        context = {
            'id_user': id_user,
            'datasoal': soal,
            'soalsoal': 'active',
            'pesan_revisi': pesan_revisi,
            'revisian': revisian
        }

        if request.user.is_authenticated:
            return render(request, 'soal-peer.html', context)
        else:
            return redirect('index')

    elif request.method == 'POST':
        form = RevisiSoalForm(request.POST)
        id_soal = soal_id
        user_id = request.user.id
        if form.is_valid():
            pesan_revisi = form.cleaned_data['pesan_revisi']
            soal = Soal.objects.get(id=id_soal)
            obj = RevisiSoal()
            obj.id_user = User.objects.get(id=user_id)
            obj.soal = soal
            obj.pesan = pesan_revisi
            obj.save()
            penulis_soal = User.objects.get(id=soal.id_creator_id)
            penugasan = Penugasan.objects.get(id=soal.indikator_id)
            soal_link = request.build_absolute_uri('/datasoal/')
            query_string = urlencode({
                'penugasan_id': soal.indikator_id
            })
            soal_link = "{}?{}".format(soal_link, query_string)
            profile_penulis = Profile.objects.get(id=penulis_soal.id)
            if profile_penulis.subskripsi_email:
                threading.Thread(target=send_writer_notif, args=(penulis_soal.email,
                                                                 penulis_soal.first_name + " " + penulis_soal.last_name,
                                                                 request.user.first_name + " " + request.user.last_name,
                                                                 pesan_revisi,
                                                                 penugasan.mata_kuliah,
                                                                 soal_link,
                                                                 request.build_absolute_uri('/subscription/'))).start()
            # status = StatusSoal.objects.get(id_soal)
            # soal.status = status
            soal.save()

            return redirect('soal-peer')

        elif request.POST.get('type') == 'lihat':
            soal = Soal.objects.get(id=id_soal)
            request.session['stem'] = soal.soal
            request.session['amount'] = soal.variasi
            request.session['arrange'] = soal.pengacakan
            request.session['var_value'] = soal.variabel
            request.session['formula'] = soal.rumus
            request.session['var_name'] = soal.var1
            request.session['option'] = soal.jawaban
            request.session['title'] = soal.judul
            request.session['img_var'] = soal.variabel_gambar
            return redirect('hasilpeer')


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def ListSoal(request):
    id_creator = request.user.id
    penugasan_id = request.session.get('penugasan_id')
    penugasan = Penugasan.objects.get(id=penugasan_id)
    soal = Soal.objects.filter(indikator=penugasan).exclude(
        id_creator=id_creator).order_by('id').reverse()
    id_user = ''
    pesan_revisi = ''
    if request.method == 'GET':
        for soals in soal:
            revisian = RevisiSoal.objects.filter(soal=soals)
            for revisi in revisian:
                id_user = revisi.id
                pesan_revisi = revisi.pesan
        context = {
            'id_user': id_user,
            'datasoal': soal,
            'soalsoal': 'active',
            'pesan_revisi': pesan_revisi,
            'revisian': revisian
        }
        if request.user.is_authenticated:
            return render(request, 'list_soal.html', context)
        else:
            return redirect('index')

    elif request.method == 'POST':
        soal_id = request.POST.get('id')
        request.session['soal_id'] = soal_id
        if 'peninjau' in request.POST:
            return redirect('assessreviewer')
        return redirect('soal-peer')


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def HasilPeer(request):
    id_creator = request.user.id
    stem = request.session.get('stem')
    amount = request.session.get('amount')
    arrange = request.session.get('arrange')
    var_value = request.session.get('var_value')
    var_name = request.session.get('var_name')
    formula = request.session.get('formula')
    option = request.session.get('option')
    img_var = request.session.get('img_var')

    if request.method == 'GET':
        # var_value = removeSpace("".join(var_value.splitlines()))
        formula = removeSpace("".join(formula.splitlines()))
        img_var = removeSpace("".join(img_var.splitlines()))

        if img_var != '---':
            var_value = var_value + ';' + img_var
        else:
            pass

        if arrange == "mapping":
            mapped_var = gen_mapped_var(var_value)

            # Check if item model contain additional formula
            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        elif arrange == "kombinasi":
            mapped_var = gen_combined_var(var_value)

            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        # format_img(mapped_var)
        pertanyaan = gen_question(stem, mapped_var)
        jawaban = gen_non_math_option(option, mapped_var)

        request.session['question'] = pertanyaan
        request.session['answer'] = jawaban
        return render(request, 'hasilpeer.html')


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def HasilPaket(request):
    if request.method == "GET":
        paket_soal_id = request.session["paket_soal_id"]
        paket_soal = PaketSoal.objects.get(id=paket_soal_id)

        jumlah_paket = paket_soal.jumlah_paket
        nama_paket = paket_soal.judul
        paket_soal_data = []
        for i in range(1, jumlah_paket+1):
            soal_paket_query = Soal_PaketSoal.objects.filter(
                paket_soal=paket_soal, nomer_paket=i).order_by("nomer_urut")
            judul = 'Paket ' + str(i)
            soal_data = []

            for soal_paket in soal_paket_query:
                soal = soal_paket.generated_soal.generated_soal
                jawaban = literal_eval(
                    soal_paket.generated_soal.generated_option)

                detail = {
                    'question': soal,
                    'option': jawaban
                }

                soal_data.append(detail)

            detail = {
                'judul': judul,
                'soal': soal_data,
                'id': i
            }

            paket_soal_data.append(detail)

    elif request.method == "POST":
        if request.POST.get('type') == 'download':
            paket_soal_id = request.session["paket_soal_id"]
            paket_soal = PaketSoal.objects.get(id=paket_soal_id)
            nomer_paket = request.POST.get('id')
            judul = paket_soal.judul + '_Paket_' + str(nomer_paket)
            soal_paket_query = Soal_PaketSoal.objects.filter(
                paket_soal=paket_soal, nomer_paket=nomer_paket).order_by("nomer_urut")

            question_data = []
            option_data = []

            for soal_paket in soal_paket_query:
                soal = soal_paket.generated_soal.generated_soal
                jawaban = literal_eval(
                    soal_paket.generated_soal.generated_option)
                question_data.append(soal)
                option_data.append(jawaban)

            request.session['question'] = question_data
            request.session['answer'] = option_data
            request.session['title'] = judul

            return redirect('test')

        elif request.POST.get('type') == 'kisi':
            paket_soal_id = request.session["paket_soal_id"]
            paket_soal = PaketSoal.objects.get(id=paket_soal_id)
            judul = paket_soal.judul
            kurikulum = paket_soal.kurikulum
            soal_brother = Soal_PaketSoal.objects.filter(
                paket_soal=paket_soal, nomer_paket=1).order_by("nomer_urut")
            soal_data = []
            for soal in soal_brother:
                no = soal.nomer_urut
                so = soal.generated_soal.parent.indikator.student_outcome
                lo = soal.generated_soal.parent.indikator.learning_outcome
                topik = soal.generated_soal.parent.indikator.topik
                matakuliah = soal.generated_soal.parent.indikator.mata_kuliah

                detail = {
                    'no': no,
                    'so': so,
                    'lo': lo,
                    'topik': topik
                }

                soal_data.append(detail)

            context = {
                'judul': judul,
                'kurikulum': kurikulum,
                'soal': soal_data,
                'matakuliah': matakuliah
            }

            template = './static/file/templatekisi.docx'
            document = download_from_template(template, context)
            document.seek(0)

            doc = document.getvalue()
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename= ' + \
                judul + '.docx'
            document.close()
            response.write(doc)

            return response

        elif request.POST.get('type') == 'kirim':
            paket_soal_id = request.session["paket_soal_id"]
            paket_soal = PaketSoal.objects.get(id=paket_soal_id)
            nomer_paket = request.POST.get('id')
            judul = paket_soal.judul + '_Paket_' + str(nomer_paket)
            soal_paket_query = Soal_PaketSoal.objects.filter(
                paket_soal=paket_soal, nomer_paket=nomer_paket).order_by("nomer_urut")

            question_data = []
            option_data = []

            for soal_paket in soal_paket_query:
                soal = soal_paket.generated_soal.generated_soal
                jawaban = literal_eval(
                    soal_paket.generated_soal.generated_option)
                question_data.append(soal)
                option_data.append(jawaban)

            request.session['question'] = question_data
            request.session['answer'] = option_data
            request.session['title'] = judul

            return redirect('kirim')

        elif request.POST.get('type') == 'send_all':
            paket_soal_id = request.session["paket_soal_id"]
            paket_soal = PaketSoal.objects.get(id=paket_soal_id)
            jumlah_paket = paket_soal.jumlah_paket

            for i in range(1, jumlah_paket+1):
                paket_soal_id = request.session["paket_soal_id"]
                paket_soal = PaketSoal.objects.get(id=paket_soal_id)
                nomer_paket = i
                judul = paket_soal.judul + '_Paket_' + str(nomer_paket)
                soal_paket_query = Soal_PaketSoal.objects.filter(
                    paket_soal=paket_soal, nomer_paket=nomer_paket).order_by("nomer_urut")

                question_data = []
                option_data = []

                for soal_paket in soal_paket_query:
                    soal = soal_paket.generated_soal.generated_soal
                    jawaban = literal_eval(
                        soal_paket.generated_soal.generated_option)
                    question_data.append(soal)
                    option_data.append(jawaban)

                request.session['question'] = question_data
                request.session['answer'] = option_data
                request.session['title'] = judul

                document = Document()

                api_upload = "https://moodle.aigyujiem.com/webservice/restful/server.php/core_files_upload"
                api_copy_private = "https://moodle.aigyujiem.com/webservice/restful/server.php/core_user_add_user_private_files"

                for soal, t in zip(request.session['question'], request.session['answer']):
                    paragraph = document.add_paragraph()
                    if 'math-tex' in soal:
                        cari = BeautifulSoup(soal).find(
                            "span", {"class": "math-tex"})
                        delete = '<span class="math-tex">' + cari.text + '</span>'
                        abcd = soal.replace(delete, ",|,")
                        final = strip_tags(unescape(abcd)).split(",")
                        a = cari.text.replace('\(', '')
                        tex_eq = a.replace('\)', '')
                        raw = str_to_raw(tex_eq)
                        eq = mcv.convert(raw)
                        tree = etree.fromstring(eq)
                        xslt = etree.parse('MML2OMML.XSL')
                        transform = etree.XSLT(xslt)
                        new_dom = transform(tree)
                        for x in range(0, len(final)):
                            if final[x] == "|":
                                paragraph._element.append(new_dom.getroot())
                            else:
                                paragraph.add_run(final[x].strip())

                    elif 'img' in soal:
                        # parsed = strip_tags(unescape(soal.strip()))
                        # paragraph.add_run(parsed)
                        cari = BeautifulSoup(soal).find('img')
                        delete = '<img alt="' + \
                            str(cari['alt']) + '" src="' + \
                            str(cari['src']) + '" />'
                        abcd = soal.replace(delete, ",|,")
                        final = strip_tags(unescape(abcd)).split(",")
                        url = BeautifulSoup(soal).find('img')['src']
                        response = requests.get(url, stream=True)
                        image = io.BytesIO(response.content)
                        image2 = Image.open(image)
                        width, height = image2.size

                        for x in range(0, len(final)):
                            if final[x] == "|":
                                if width >= 288 and width >= height:
                                    paragraph.add_run().add_picture(image, width=Inches(3.0))
                                elif height >= 192 and height > width:
                                    paragraph.add_run().add_picture(image, height=Inches(2.0))
                                else:
                                    paragraph.add_run().add_picture(image)
                            else:
                                paragraph.add_run(final[x].strip())

                    else:
                        parsed = strip_tags(unescape(soal.strip()))
                        paragraph.add_run(parsed)
                    paragraph = document.add_paragraph()
                    for j in t:
                        if 'img' in j:
                            cari = BeautifulSoup(j).find('img')
                            delete = '<img alt="' + \
                                str(cari['alt']) + '" src="' + \
                                str(cari['src']) + '" />'
                            abcd = j.replace(delete, ",|,")
                            final = strip_tags(unescape(abcd)).split(",")
                            url = BeautifulSoup(j).find('img')['src']
                            response = requests.get(url, stream=True)
                            image = io.BytesIO(response.content)
                            image2 = Image.open(image)
                            width, height = image2.size

                            for x in range(0, len(final)):
                                if final[x] == "|":
                                    if width >= 288 and width >= height:
                                        document.add_picture(
                                            image, width=Inches(3.0))
                                        document.add_paragraph()
                                    elif height >= 192 and height > width:
                                        document.add_picture(
                                            image, height=Inches(2.0))
                                        document.add_paragraph()
                                    else:
                                        document.add_picture(image)
                                        document.add_paragraph()
                                else:
                                    paragraph.add_run(final[x].strip())
                        else:
                            if '\(' in j:
                                y = j.replace('\(', '<span class="math-tex">')
                                z = y.replace('\)', '</span>')
                                cari2 = BeautifulSoup(z).find(
                                    "span", {"class": "math-tex"})
                                delete = '<span class="math-tex">' + cari2.text + '</span>'
                                abcd = unescape(z).replace(delete, ",|,")
                                final = abcd.split(",")
                                tex_eq = cari2.text
                                raw = str_to_raw(tex_eq)
                                eq = mcv.convert(raw)
                                tree = etree.fromstring(eq)
                                xslt = etree.parse('MML2OMML.XSL')
                                transform = etree.XSLT(xslt)
                                new_dom = transform(tree)
                                for x in range(0, len(final)):
                                    if final[x] == "|":
                                        paragraph._element.append(
                                            new_dom.getroot())
                                    else:
                                        paragraph.add_run(final[x].strip())
                                paragraph = document.add_paragraph()
                            else:
                                parsed = strip_tags(unescape(j.strip()))
                                paragraph.add_run(parsed)
                                paragraph = document.add_paragraph()

                response = HttpResponse(
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename= ' + \
                    request.session['title'] + '.docx'
                file_stream = io.BytesIO()
                document.save(response)
                document.save(file_stream)
                file_bytes = file_stream.getvalue()
                base64_encoded_data = base64.b64encode(file_bytes)
                base64_message = base64_encoded_data.decode(('utf-8'))

                pathname = str(request.session['title'].split('_'))

                header = {
                    'Authorization': request.session['moodle_token'],
                    'Accept': 'application/json',
                    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.97 Safari/537.36"
                }

                payload_send = {
                    "component": "user",
                    "filearea": "draft",
                    "itemid": 0,
                    "filepath": "/",
                    "filename": request.session['title'] + '.docx',
                    "contextlevel": "user",
                    "instanceid": request.session['user_id'],
                    "filecontent": base64_message
                }

                moodle_response_send = requests.post(
                    api_upload, headers=header, data=payload_send)
                request.session['file_id'] = moodle_response_send.json()[
                    'itemid']

                payload_copy = {
                    "draftid": request.session['file_id']
                }

                moodle_response_copy = requests.post(
                    api_copy_private, headers=header, data=payload_copy)
        return redirect('hasilpaket')

    return render(request, 'hasilpaket.html', {'paket_soal_data': paket_soal_data, 'judul': nama_paket})


@csrf_exempt
def show_pdf(request):
    dir_list = {
        'fisika-teknik': ['uas_gasal2018', 'uas_gasal2019', 'uas_genap2017', 'uas_genap2018', 'uas_genap2019',
                          'uas_genap2020', 'uts_gasal2018', 'uts_gasal2019', 'uts_genap2018', 'uts_genap2020'
                          ]
    }
    file_dir = 'quiz/similaritychecker/data'
    if (request.method == 'GET') and request.user.is_authenticated:
        matkul = str(request.GET.get('matkul'))
        if matkul in dir_list:
            file_name = str(request.GET.get('file'))
            if file_name in dir_list[matkul]:
                filename = matkul + '_' + file_name
                new_file_dir = file_dir + '/' + matkul + '/' + file_name + '.pdf'
                return FileResponse(open(new_file_dir, 'rb'), filename=filename)

    return Http404("File not exists!")


@csrf_exempt
def subscription(request):
    if not request.user.is_authenticated:
        base_url = reverse('index')
        query_string = urlencode({
            'url': 'subscription'
        })
        url = '{}?{}'.format(base_url, query_string)
        return redirect(url)

    user_email = request.user.email
    user_profile = Profile.objects.get(id=request.user.id)
    user_subskripsi = user_profile.subskripsi_email
    if request.method == 'GET':
        return render(request, 'subscription.html', {
            'email': user_email,
            'subskripsi': user_subskripsi
        })
    elif request.method == 'POST':
        hasil_subskripsi = request.POST.get('subskripsi')
        if hasil_subskripsi == 'True':
            hasil_subskripsi = True
        else:
            hasil_subskripsi = False
        user_profile.subskripsi_email = hasil_subskripsi
        user_profile.save()
        messages.success(request, "Konfigurasi langganan berhasil diubah!")
        return render(request, 'subscription.html', {
            'email': user_email,
            'subskripsi': hasil_subskripsi
        })
    return Http404()


@user_passes_test(lambda user: Group.objects.get(name='akademik') in user.groups.all())
def assess_reviewer(request):
    paket_soal_id = request.session.get('paket_soal_id')

    paket_soal = PaketSoal.objects.get(id=paket_soal_id)
    creator_group = Group.objects.get(name='creator')
    users_creator = User.objects.filter(groups=creator_group)
    creator = paket_soal.creator
    context = {
        'soal': paket_soal,
        'creator': creator,
        'users': users_creator
    }
    if request.method == 'POST':
        if "kembali" in request.POST:
            return redirect('daftarpaket')
        reviewer_id = request.POST.get('nama_peninjau')

        if reviewer_id != str(0):
            user_peninjau_1 = User.objects.get(id=reviewer_id)
            email_peninjau_1 = user_peninjau_1.email
            test_peninjau_1 = Peninjau.objects.filter(
                peninjau_user_id=user_peninjau_1.id, paket_soal=paket_soal)
            if not test_peninjau_1.exists():
                status = StatusPaketSoal.objects.get(id=1)
                peninjau_1 = Peninjau(
                    peninjau_user=user_peninjau_1,
                    paket_soal=paket_soal,
                    penunjuk_user=request.user,
                    tanggal_penunjukkan=datetime.now(),
                    status=status
                )
                peninjau_1.save()

                profile_peninjau = Profile.objects.get(user=user_peninjau_1)
                if profile_peninjau.subskripsi_email:
                    threading.Thread(target=send_reviewer_notif, args=(
                        email_peninjau_1,
                        f"{user_peninjau_1.first_name} {user_peninjau_1.last_name}",
                        f"{request.user.first_name} {request.user.last_name}",
                        paket_soal.judul,
                        f"{paket_soal.creator.first_name} {paket_soal.creator.last_name}",
                        f"{paket_soal.tahun_ajaran} {paket_soal.semester}",
                        request.build_absolute_uri('/soal-peninjau/'),
                        request.build_absolute_uri('/subscription/')
                    )).start()

                soal_paketsoals = Soal_PaketSoal.objects.filter(
                    paket_soal=paket_soal, nomer_paket=1)
                for soal in soal_paketsoals:
                    hasil_review = HasilReview(
                        status=status,
                        peninjau=peninjau_1,
                        soal_paketSoal=soal
                    )
                    hasil_review.save()
                # profile_peninjau_1 = Profile.objects.get(id = user_peninjau_1.id)
                # if profile_peninjau_1.subskripsi_email:
                    # threading.Thread(target=send_reviewer_notif, args=(email_peninjau_1,
                    #     user_peninjau_1.first_name + " " + user_peninjau_1.last_name,
                    #     request.user.first_name + " " + request.user.last_name,
                    #     soal.judul,
                    #     penulis_soal.first_name + " " + penulis_soal.last_name,
                    #     penugasan.mata_kuliah,
                    #     soal_link,
                    #     request.build_absolute_uri('/soal-peninjau/'),
                    #     request.build_absolute_uri('/subscription/'))).start()
                messages.success(
                    request, "Pengguna berhasil ditunjuk sebagai reviewer")
            else:
                messages.error(
                    request, "Pengguna sudah ditunjuk menjadi reviewer!")
        else:
            messages.error(request, "Silahkan pilih Pengguna!")
    return render(request, 'assess-reviewer.html', context)


@csrf_exempt
@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def soal_peninjau(request):
    if not request.user.is_authenticated:
        base_url = reverse('index')
        query_string = urlencode({
            'url': 'soalpeninjau'
        })
        url = '{}?{}'.format(base_url, query_string)
        return redirect(url)

    peninjaus = Peninjau.objects.filter(
        peninjau_user_id=request.user.id).order_by('id').reverse()
    context = {
        'peninjaus': peninjaus
    }
    if request.method == 'POST':
        if request.POST.get('type') == 'review':
            reviewer_id = request.POST.get('id')
            request.session['reviewer_id'] = reviewer_id

            return redirect('soalpaketpeninjau')
        # elif 'peninjau' in request.POST:
        #     peninjau_id = request.POST.get('peninjau_id')
        #     peninjau = Peninjau.objects.get(id=peninjau_id)
        #     peninjau.delete()
        #     messages.success(request, "Status berhasil dihapus!")

    return render(request, 'soal-peninjau.html', context)


@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def soalPaket_peninjau(request):
    reviewer_id = request.session.get('reviewer_id')
    reviewer = Peninjau.objects.get(id=reviewer_id)
    soals = Soal_PaketSoal.objects.filter(
        paket_soal=reviewer.paket_soal, nomer_paket=1)
    jumlah_soal = soals.count()
    hasil_reviews = []
    for soal in soals:
        parent_soal = soal.generated_soal.parent
        indikator = parent_soal.indikator
        hasil = HasilReview.objects.get(peninjau=reviewer, soal_paketSoal=soal)

        detail = {
            'parent_soal': parent_soal,
            'indikator': indikator,
            'hasil': hasil
        }
        hasil_reviews.append(detail)
    context = {
        'soals': hasil_reviews,
        'paket_soal': reviewer.paket_soal,
        'reviewer': reviewer,
        'jumlah_soal': jumlah_soal
    }
    if request.method == 'POST':
        if request.POST.get('type') == 'review':
            hasil_id = request.POST.get('id')
            print(hasil_id)
            request.session['hasil_id'] = hasil_id

            return redirect('berireview')

    return render(request, 'soal-paket-review.html', context)


@user_passes_test(lambda user: Group.objects.get(name='creator') in user.groups.all())
def beriReview(request):
    hasil_id = request.session.get('hasil_id')
    hasil_review = HasilReview.objects.get(id=hasil_id)
    soalpaket = hasil_review.soal_paketSoal
    paket = soalpaket.paket_soal
    soalpakets = Soal_PaketSoal.objects.filter(
        paket_soal=paket, nomer_urut=soalpaket.nomer_urut)
    max_soal = 1
    list_pertanyaan = []
    list_jawaban = []
    for soal in soalpakets:
        list_pertanyaan.append(soal.generated_soal.generated_soal)
        list_jawaban.append(soal.generated_soal.generated_option)
        if soal.nomer_paket > max_soal:
            max_soal = soal.nomer_paket
    initial = {
        'kejelasan_bahasa': hasil_review.kejelasan_bahasa,
        'kejelasan_gambar': hasil_review.kejelasan_gambar,
        'kesesuaian_waktu': hasil_review.kesesuaian_waktu,
        'kesesuaian_cpmk': hasil_review.kesesuaian_cpmk,
        'komentar': hasil_review.komentar,
        'status': hasil_review.status.id,
    }
    form = BeriReviewForm(data=initial)
    if request.method == 'POST':
        form = BeriReviewForm(request.POST)
        if form.is_valid():
            hasil_review.kejelasan_bahasa = form.cleaned_data['kejelasan_bahasa']
            hasil_review.kejelasan_gambar = form.cleaned_data['kejelasan_gambar']
            hasil_review.kesesuaian_waktu = form.cleaned_data['kesesuaian_waktu']
            hasil_review.kesesuaian_cpmk = form.cleaned_data['kesesuaian_cpmk']
            hasil_review.komentar = form.cleaned_data['komentar']
            hasil_review.status = StatusPaketSoal.objects.get(
                id=form.cleaned_data['status'])
            hasil_review.created_date = datetime.now()
            hasil_review.save()
            status_paket = 3
            reviewer_id = request.session.get('reviewer_id')
            reviewer = Peninjau.objects.get(id=reviewer_id)
            hasil_reviews = HasilReview.objects.filter(peninjau=reviewer)
            for hasil in hasil_reviews:
                if (hasil.status.id < status_paket):
                    status_paket = hasil.status.id
            status_reviewer = StatusPaketSoal.objects.get(id=status_paket)
            reviewer.status = status_reviewer
            reviewer.save()
            status_soalpaket = 3
            hasil_reviews = HasilReview.objects.filter(
                soal_paketSoal=soalpaket)
            for hasil in hasil_reviews:
                if hasil.status.id < status_soalpaket:
                    status_soalpaket = hasil.status.id
            soalpaket.status = StatusPaketSoal.objects.get(id=status_soalpaket)
            soalpaket.save()
            status_paketsoal = 3
            paketsoals = Soal_PaketSoal.objects.filter(
                paket_soal=paket, nomer_paket=1)
            for paketsoal in paketsoals:
                if (paketsoal.status.id < status_paketsoal):
                    status_paketsoal = paketsoal.status.id
            if status_paketsoal == 1:
                status_paketsoal = 2
            status_baru = StatusPaketSoal.objects.get(id=status_paketsoal)
            paket.status = status_baru
            paket.save()
            profile_penulis = Profile.objects.get(user=paket.creator)
            if profile_penulis.subskripsi_email:
                threading.Thread(target=send_penulis_notif_reviewed, args=(
                    paket.creator.email,
                    f"{paket.creator.first_name} {paket.creator.last_name}",
                    f"{request.user.first_name} {request.user.last_name}",
                    paket.judul,
                    soalpaket.nomer_urut,
                    request.build_absolute_uri('/datapaket/'),
                    request.build_absolute_uri('/subscription/')
                )).start()

            messages.success(request, "Review berhasil diperbarui")
    question = soalpaket.generated_soal.generated_soal
    option = soalpaket.generated_soal.generated_option
    so = soalpaket.generated_soal.parent.indikator.student_outcome
    lo = soalpaket.generated_soal.parent.indikator.learning_outcome

    detail = {
        'question': question,
        'option': literal_eval(option),
        'nomor': soalpaket.nomer_paket,
        'so': so,
        'lo': lo,
        'max': max_soal,
        'pertanyaan': list_pertanyaan,
        'jawaban': list_jawaban
    }

    return render(request, 'berireview.html', {'forms': form, 'detail': detail})


@user_passes_test(lambda user: Group.objects.get(name='akademik') in user.groups.all())
def DaftarPaket(request):
    if request.method == 'GET':
        paket_soal_query = PaketSoal.objects.all().order_by('id')
        paket_soal_data = []

        for paket_soal in paket_soal_query:
            judul = paket_soal.judul

            created_date = paket_soal.created_date
            kurikulum = paket_soal.kurikulum
            status = paket_soal.status.id

            id = paket_soal.id

            detail = {
                'judul': judul,
                'id': id,
                'created_date': created_date,
                'kurikulum': kurikulum,
                'semester': str(paket_soal.tahun_ajaran) + " " + paket_soal.semester,
                'status': status
            }

            paket_soal_data.append(detail)

    elif request.method == 'POST':
        paket_soal_id = request.POST.get('id')

        if request.POST.get('type') == 'review':
            request.session['paket_soal_id'] = paket_soal_id

            return redirect('assessreviewer')
        elif request.POST.get('type') == 'daftar':
            request.session['paket_soal_id'] = paket_soal_id

            return redirect('daftarreviewer')

    return render(request, 'daftarpaket.html', {'paket_soal': paket_soal_data, 'datapaket': 'active'})


@user_passes_test(lambda user: Group.objects.get(name='akademik') in user.groups.all())
def DaftarReviewer(request):
    paket_soal_id = request.session.get('paket_soal_id')
    paket_soal = PaketSoal.objects.get(id=paket_soal_id)
    reviewers = Peninjau.objects.filter(paket_soal=paket_soal)
    context = {
        'reviewers': reviewers,
        'judul': paket_soal.judul
    }
    if request.method == 'POST':
        if request.POST.get('type') == 'hapus':
            reviewer_id = request.POST.get('id')
            reviewer = Peninjau.objects.get(id=reviewer_id)
            soals = Soal_PaketSoal.objects.filter(
                paket_soal=paket_soal, nomer_paket=1)
            for soal in soals:
                hasil_reviews = HasilReview.objects.get(
                    peninjau=reviewer, soal_paketSoal=soal)
                hasil_reviews.delete()
            reviewer.delete()
            messages.success(request, "Reviewer berhasil dihapus")
    return render(request, 'daftarreviewer.html', context=context)


@csrf_exempt
def UploadImg(request):
    # config = {
    #     "apiKey": "AIzaSyDcD37EN-yzuIYkt0SuFVFYrMfa1g8XL20",
    #     "authDomain": "aigyujiem-1414b.firebaseapp.com",
    #     "projectId": "aigyujiem-1414b",
    #     "storageBucket": "aigyujiem-1414b.appspot.com",
    #     "messagingSenderId": "957774555557",
    #     "appId": "1:957774555557:web:a6249e2d2d5c8fd2c7a04c",
    #     "measurementId": "G-ZSPBZD32FT",
    #     "databaseURL": "https://aigyujiem-1414b-default-rtdb.asia-southeast1.firebasedatabase.app/"
    # }

    # firebase = pyrebase.initialize_app(config)
    # storage = firebase.storage()
    if request.method == "POST":
        file_obj = request.FILES['file']
        file_name_suffix = file_obj.name.split(".")[-1]
        if file_name_suffix not in ["jpg", "png", "gif", "jpeg", ]:
            return JsonResponse({"message": "Wrong file format"})

        path = os.path.join(
            MEDIA_ROOT,
            'tinymce',
        )
        # If there is no such path, create
        if not os.path.exists(path):
            os.makedirs(path)

        file_path = os.path.join(path, file_obj.name)

        file_url = f'{MEDIA_URL}tinymce/{file_obj.name}'

        if os.path.exists(file_path):
            return JsonResponse({
                "message": "file already exist",
                'location': file_url
            })

        with open(file_path, 'wb+') as f:
            for chunk in file_obj.chunks():
                f.write(chunk)

        return JsonResponse({
            'message': 'Image uploaded successfully',
            'location': file_url
        })
    return JsonResponse({'detail': "Wrong request"})
    # if request.method == 'POST':
    #     file = request.FILES['file']
    #     file_name_suffix = file.name.split(".")[-1]
    #     if file_name_suffix not in ["jpg", "png", "gif", "jpeg", ]:
    #         return JsonResponse({"message": "Wrong file format"})
    #     file_save = default_storage.save(file.name, file)
    #     storage.child("files/" + file.name).put("media/" + file.name)
    #     delete = default_storage.delete(file.name)
    #     messages.success(request, "File upload in Firebase Storage successful")
    #     return redirect('identitas')
    # else:
    #     return render(request, 'identitas.html')


def render_pdf_view(request):
    id_creator = request.user.id
    id_soal = request.session.get('id_soal')
    stem = request.session.get('stem')
    amount = request.session.get('amount')
    arrange = request.session.get('arrange')
    var_value = request.session.get('var_value')
    var_name = request.session.get('var_name')
    formula = request.session.get('formula')
    option = request.session.get('option')
    img_var = request.session.get('img_var')
    # soal = Soal.objects.get(id=id_soal)
    # status = soal.status.id

    if request.method == 'GET':
        # var_value = removeSpace("".join(var_value.splitlines()))
        formula = removeSpace("".join(formula.splitlines()))
        if img_var:
            img_var = removeSpace("".join(img_var.splitlines()))

            if img_var != '---':
                var_value = var_value + ';' + img_var
            else:
                pass

        if arrange == "mapping":
            mapped_var = gen_mapped_var(var_value)

            # Check if item model contain additional formula
            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        elif arrange == "kombinasi":
            mapped_var = gen_combined_var(var_value)

            if formula != '---':
                update_var(mapped_var, formula)
            else:
                pass

        # format_img(mapped_var)
        pertanyaan = gen_question(stem, mapped_var)
        jawaban = gen_non_math_option(option, mapped_var)

        request.session['question'] = pertanyaan
        request.session['answer'] = jawaban

        soal = Soal.objects.get(id=id_soal)
        template_path = 'renderpdf.html'
        context = {'pertanyaan': pertanyaan,
                   'jawaban': jawaban
                   }
        # Create a Django response object, and specify content_type as pdf
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename=' + \
            request.session['title'] + '.pdf'
        # find the template and render it.
        template = get_template(template_path)
        html = template.render(context)

        # create a pdf
        pisa_status = pisa.CreatePDF(
            html, dest=response)
        # if error then show some funny view
        if pisa_status.err:
            return HttpResponse('We had some errors <pre>' + html + '</pre>')
        return response


def cobaDulu(request):
    return render(request, 'coba.html')
